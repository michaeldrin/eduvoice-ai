import os
import logging
import uuid
import re
import random
from datetime import datetime
import json
from functools import wraps
from io import BytesIO
from traceback import format_exc

# Define placeholder classes for libraries that might not be installed
class FitzPlaceholder:
    def open(self, *args, **kwargs):
        raise ImportError("PyMuPDF (fitz) module is not installed")

class DocxPlaceholder:
    def Document(self, *args, **kwargs):
        raise ImportError("python-docx module is not installed")

class PILPlaceholder:
    def open(self, *args, **kwargs):
        raise ImportError("PIL/Pillow module is not installed")

class PytesseractPlaceholder:
    def image_to_string(self, *args, **kwargs):
        raise ImportError("pytesseract module is not installed")

# Import required libraries with fallbacks
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = FitzPlaceholder()

try:
    import docx
except ImportError:
    docx = DocxPlaceholder()

try:
    from PIL import Image
except ImportError:
    Image = PILPlaceholder()

try:
    import pytesseract
except ImportError:
    pytesseract = PytesseractPlaceholder()

from flask import Flask, render_template, request, redirect, url_for, jsonify, session
from flask import send_from_directory, flash
from werkzeug.utils import secure_filename
from openai import OpenAI
from models import db, Document, ChatMessage

# Set up logging
os.makedirs('logs', exist_ok=True)
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev_secret_key")

# Custom Jinja2 filter for converting newlines to <br>
@app.template_filter('nl2br')
def nl2br(value):
    if value:
        return value.replace('\n', '<br>')
    return value

# Add global template variables
@app.context_processor
def inject_global_variables():
    return {
        'current_year': datetime.now().year
    }

# Configure upload folder and allowed extensions
UPLOAD_FOLDER = 'uploads'
ATTACHMENT_FOLDER = os.path.join(UPLOAD_FOLDER, 'attachments')
# Define file extension groups
PDF_EXTENSIONS = {'pdf'}
DOCX_EXTENSIONS = {'docx'}
TEXT_EXTENSIONS = {'txt'}
IMAGE_EXTENSIONS = {'jpg', 'jpeg', 'png'}
ALLOWED_EXTENSIONS = PDF_EXTENSIONS.union(DOCX_EXTENSIONS).union(TEXT_EXTENSIONS).union(IMAGE_EXTENSIONS)

# Define supported OCR languages
OCR_LANGUAGES = {
    'eng': 'English',
    'deu': 'German',
    'fra': 'French',
    'spa': 'Spanish',
    'ara': 'Arabic',
    'fas': 'Persian',
    'urd': 'Urdu',
    'rus': 'Russian',
    'ukr': 'Ukrainian',
    'chi_sim': 'Chinese (Simplified)',
    'chi_tra': 'Chinese (Traditional)',
    'jpn': 'Japanese',
    'kor': 'Korean'
}

# Define supported translation languages with flags and native names
TRANSLATION_LANGUAGES = {
    'es': {'name': 'Spanish', 'native': 'Español', 'flag': '🇪🇸'},
    'fr': {'name': 'French', 'native': 'Français', 'flag': '🇫🇷'},
    'de': {'name': 'German', 'native': 'Deutsch', 'flag': '🇩🇪'},
    'it': {'name': 'Italian', 'native': 'Italiano', 'flag': '🇮🇹'},
    'pt': {'name': 'Portuguese', 'native': 'Português', 'flag': '🇵🇹'},
    'ru': {'name': 'Russian', 'native': 'Русский', 'flag': '🇷🇺'},
    'zh': {'name': 'Chinese (Simplified)', 'native': '中文', 'flag': '🇨🇳'},
    'ja': {'name': 'Japanese', 'native': '日本語', 'flag': '🇯🇵'},
    'ko': {'name': 'Korean', 'native': '한국어', 'flag': '🇰🇷'},
    'ar': {'name': 'Arabic', 'native': 'العربية', 'flag': '🇸🇦'},
    'hi': {'name': 'Hindi', 'native': 'हिन्दी', 'flag': '🇮🇳'},
    'en': {'name': 'English', 'native': 'English', 'flag': '🇺🇸'}
}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['ATTACHMENT_FOLDER'] = ATTACHMENT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

# Configure database
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///app.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

# Create necessary directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ATTACHMENT_FOLDER, exist_ok=True)
os.makedirs('static/summaries', exist_ok=True)

# Database retry decorator
def db_retry(max_retries=3):
    """Decorator to retry database operations on connection errors"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            retries = 0
            while retries < max_retries:
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    retries += 1
                    logger.warning(f"Database operation failed (attempt {retries}/{max_retries}): {str(e)}")
                    
                    # Check if it's a connection error we can retry
                    error_str = str(e).lower()
                    if "operational error" in error_str or "connection" in error_str:
                        # Do a rollback to reset the session state
                        try:
                            db.session.rollback()
                            logger.info("Session rolled back successfully")
                        except Exception as rollback_error:
                            logger.error(f"Failed to rollback session: {rollback_error}")
                            
                        # Last attempt, let it fail normally
                        if retries >= max_retries:
                            logger.error(f"Maximum retries reached ({max_retries}), operation failed")
                            raise
                    else:
                        # Not a connection error we can retry, re-raise immediately
                        raise
            # This should not be reached normally, but just in case
            raise Exception(f"Database operation failed after {max_retries} retries")
        return wrapper
    return decorator

# Initialize database
with app.app_context():
    try:
        # Create tables (they should already exist from SQL initialization)
        db.create_all()
        logger.info("Database tables initialized successfully")
    except Exception as e:
        logger.error(f"Error initializing database: {e}")
        # We'll continue and let individual requests handle specific errors

# Helper function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper function to check if a file is an image
def is_image_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in IMAGE_EXTENSIONS

# Helper function to extract text from image files using OCR
def extract_text_from_image(file_path, language='eng'):
    """
    Extract text from image file using pytesseract OCR
    
    Args:
        file_path: Path to the image file
        language: OCR language. Defaults to 'eng' (English).
                  Can be 'eng' (English), 'deu' (German), 'fas' (Persian),
                  'fra' (French), 'spa' (Spanish), etc. or combined with '+' 
                  like 'eng+deu' for multiple languages.
    
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"Image file not found: {file_path}")
            return None, "Image file not found. Please upload the file again."
        
        # Check if pytesseract is properly installed
        if isinstance(pytesseract, PytesseractPlaceholder):
            logger.error("pytesseract module is not installed")
            return None, "Image OCR is currently unavailable. The server is missing required libraries."
        
        # Open the image with PIL
        image = Image.open(file_path)
        
        # Try to detect script/language if not specified
        # First try with default English
        text = pytesseract.image_to_string(image, lang='eng')
        
        # If English doesn't give good results, try other common languages
        if not text or len(text.strip()) < 10:  # Very little text was found
            logger.info(f"Trying multi-language OCR for image: {file_path}")
            
            # Try common languages one by one
            for lang_code in ['eng+deu+fra+spa', 'ara+fas+urd', 'rus+ukr+bel', 'chi_sim+chi_tra+jpn+kor']:
                try:
                    lang_text = pytesseract.image_to_string(image, lang=lang_code)
                    if lang_text and len(lang_text.strip()) > len(text.strip()):
                        text = lang_text
                        logger.info(f"Better OCR results with language(s): {lang_code}")
                except Exception as lang_err:
                    logger.warning(f"Error trying language {lang_code}: {lang_err}")
                    continue
        
        # Preprocess the image if text extraction was poor
        if not text or len(text.strip()) < 10:
            logger.info(f"Trying image preprocessing for better OCR: {file_path}")
            try:
                # Convert to grayscale
                gray_image = image.convert('L')
                
                # Apply thresholding to improve contrast
                threshold_image = gray_image.point(lambda x: 0 if x < 128 else 255, '1')
                
                # Try OCR again on the processed image
                processed_text = pytesseract.image_to_string(threshold_image)
                
                if processed_text and len(processed_text.strip()) > len(text.strip()):
                    text = processed_text
                    logger.info("Image preprocessing improved OCR results")
            except Exception as process_err:
                logger.warning(f"Error during image preprocessing: {process_err}")
        
        # Check final result
        if not text or len(text.strip()) == 0:
            logger.warning(f"No text extracted from image: {file_path}")
            return None, "No readable text found in the image file. The image may not contain clear text."
            
        # Add info about multi-language support
        info_text = "This text was extracted using OCR (Optical Character Recognition) technology with multi-language support.\n\n"
        return info_text + text.strip(), None
        
    except Exception as e:
        logger.exception(f"Error extracting text from image: {e}")
        return None, f"Error processing image file: {str(e)}"

# Helper function to extract text from plain text files
def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"Text file not found: {file_path}")
            return None, "Text file not found. Please upload the file again."
        
        # Open and read the text file
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        
        # Check if we got any text
        if not text or len(text.strip()) == 0:
            logger.warning(f"No text extracted from text file: {file_path}")
            return None, "The text file appears to be empty."
            
        return text, None
        
    except Exception as e:
        logger.exception(f"Error extracting text from text file: {e}")
        return None, f"Error processing text file: {str(e)}"

# Helper function to extract text from PDF files
def extract_text_from_pdf(file_path):
    """
    Extract text, images, and tables from PDF file
    
    Uses PyMuPDF for basic text and image extraction,
    pdfplumber for table extraction, and pytesseract for OCR on images
    """
    try:
        text = ""
        image_texts = []
        table_texts = []
        
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"PDF file not found: {file_path}")
            return None, "PDF file not found. Please upload the file again."
        
        # Check if PyMuPDF is properly installed
        if isinstance(fitz, FitzPlaceholder):
            logger.error("PyMuPDF (fitz) module is not installed")
            return None, "PDF processing is currently unavailable. The server is missing required libraries."
         
        # 1. Extract text and images using PyMuPDF
        with fitz.open(file_path) as pdf_document:
            if len(pdf_document) == 0:
                logger.warning(f"PDF has no pages: {file_path}")
                return None, "The PDF file doesn't contain any pages."
                
            # Extract text from each page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_text = page.get_text()
                text += page_text
                
                # Extract images from this page
                image_list = page.get_images(full=True)
                
                # Process each image on the page
                for img_index, img_info in enumerate(image_list):
                    try:
                        xref = img_info[0]  # Get reference number for the image
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Need to save the image temporarily for OCR
                        temp_img_path = os.path.join(
                            os.path.dirname(file_path),
                            f"temp_img_p{page_num+1}_{img_index}.jpg"
                        )
                        
                        with open(temp_img_path, "wb") as temp_file:
                            temp_file.write(image_bytes)
                        
                        # Check image size - skip tiny images (likely icons)
                        img = Image.open(temp_img_path)
                        width, height = img.size
                        
                        # Skip small images (likely icons or decorative elements)
                        if width < 100 or height < 100:
                            os.remove(temp_img_path)
                            continue
                        
                        # Process with OCR if pytesseract is available
                        if not isinstance(pytesseract, PytesseractPlaceholder):
                            # Try OCR on the image
                            ocr_text = pytesseract.image_to_string(img)
                            
                            # Only add if meaningful text was found (more than 10 chars)
                            if ocr_text and len(ocr_text.strip()) > 10:
                                image_texts.append({
                                    'page': page_num + 1,
                                    'text': ocr_text.strip()
                                })
                                logger.info(f"Extracted OCR text from image on page {page_num+1}")
                        
                        # Clean up the temporary file
                        os.remove(temp_img_path)
                        
                    except Exception as img_err:
                        logger.warning(f"Error processing image on page {page_num+1}: {str(img_err)}")
                        # Continue with next image, don't stop the whole process
        
        # 2. Extract tables using pdfplumber and camelot
        try:
            # First try with pdfplumber which is easier to use but less powerful
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    
                    if tables:
                        for table_idx, table in enumerate(tables):
                            # Convert table to text format
                            if table and len(table) > 0:
                                table_text = ""
                                
                                # Process each row
                                for row in table:
                                    # Filter None values and convert to strings
                                    processed_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                    # Add row text
                                    table_text += " | ".join(processed_row) + "\n"
                                
                                if table_text.strip():
                                    table_texts.append({
                                        'page': page_num + 1,
                                        'text': table_text.strip()
                                    })
                                    logger.info(f"Extracted table {table_idx+1} from page {page_num+1} using pdfplumber")
            
            # Then try with camelot for more complex tables
            try:
                import camelot
                
                # Camelot requires "lattice" or "stream" mode
                # "lattice" is for tables with visible borders
                # "stream" is for tables without clear borders
                camelot_tables = camelot.read_pdf(file_path, pages='all', flavor='lattice')
                
                # Check if any tables were found
                if camelot_tables and len(camelot_tables) > 0:
                    for i, table in enumerate(camelot_tables):
                        # Get the table page number
                        table_page = table.page
                        
                        # Convert to dataframe then to string
                        df = table.df
                        
                        # Skip tables that are too small (likely false positives)
                        if df.shape[0] <= 1 or df.shape[1] <= 1:
                            continue
                            
                        # Format the table as text
                        table_text = ""
                        for _, row in df.iterrows():
                            # Join cells with separator, skip empty cells
                            row_text = " | ".join([str(cell).strip() for cell in row if str(cell).strip()])
                            if row_text:
                                table_text += row_text + "\n"
                        
                        # Check if this table provides new information
                        # Only add if it has meaningful content
                        if table_text.strip():
                            # Check if this table is already extracted by pdfplumber
                            is_duplicate = False
                            for existing_table in table_texts:
                                if existing_table['page'] == int(table_page):
                                    # Simple check for similarity - if more than 70% of lines are identical
                                    existing_lines = set(existing_table['text'].split('\n'))
                                    new_lines = set(table_text.split('\n'))
                                    common_lines = existing_lines.intersection(new_lines)
                                    
                                    if len(common_lines) / max(len(existing_lines), len(new_lines)) > 0.7:
                                        is_duplicate = True
                                        break
                            
                            if not is_duplicate:
                                table_texts.append({
                                    'page': int(table_page),
                                    'text': table_text.strip()
                                })
                                logger.info(f"Extracted additional table from page {table_page} using camelot")
                
                # Try with stream mode for tables without borders
                try:
                    stream_tables = camelot.read_pdf(file_path, pages='all', flavor='stream')
                    
                    if stream_tables and len(stream_tables) > 0:
                        for i, table in enumerate(stream_tables):
                            # Get the table page number
                            table_page = table.page
                            
                            # Convert to dataframe then to string
                            df = table.df
                            
                            # Skip tables that are too small (likely false positives)
                            if df.shape[0] <= 1 or df.shape[1] <= 1:
                                continue
                                
                            # Format the table as text
                            table_text = ""
                            for _, row in df.iterrows():
                                # Join cells with separator, skip empty cells
                                row_text = " | ".join([str(cell).strip() for cell in row if str(cell).strip()])
                                if row_text:
                                    table_text += row_text + "\n"
                            
                            # Only add if it has meaningful content and is not a duplicate
                            if table_text.strip():
                                # Check for duplicates
                                is_duplicate = False
                                for existing_table in table_texts:
                                    if existing_table['page'] == int(table_page):
                                        existing_lines = set(existing_table['text'].split('\n'))
                                        new_lines = set(table_text.split('\n'))
                                        common_lines = existing_lines.intersection(new_lines)
                                        
                                        if len(common_lines) / max(len(existing_lines), len(new_lines)) > 0.7:
                                            is_duplicate = True
                                            break
                                
                                if not is_duplicate:
                                    table_texts.append({
                                        'page': int(table_page),
                                        'text': table_text.strip()
                                    })
                                    logger.info(f"Extracted stream table from page {table_page} using camelot")
                    
                except Exception as stream_err:
                    logger.warning(f"Error with camelot stream mode: {str(stream_err)}")
                
            except ImportError:
                logger.warning("camelot-py not available for advanced table extraction")
            except Exception as camelot_err:
                logger.warning(f"Error using camelot for table extraction: {str(camelot_err)}")
                
        except ImportError:
            logger.warning("pdfplumber not available for table extraction")
        except Exception as table_err:
            logger.warning(f"Error extracting tables: {str(table_err)}")
        
        # 3. Combine all extracted content
        combined_text = text
        
        # Add image texts if any were found
        for img_data in image_texts:
            combined_text += f"\n\n[Extracted from image on page {img_data['page']}]\n{img_data['text']}"
        
        # Add table texts if any were found
        for table_data in table_texts:
            combined_text += f"\n\n[Extracted from table on page {table_data['page']}]\n{table_data['text']}"
        
        # Check if we got any text
        if not combined_text or len(combined_text.strip()) == 0:
            logger.warning(f"No text extracted from PDF: {file_path}")
            return None, "No readable text found in the PDF file. It may contain only images or be password protected."
        
        # Log extraction stats
        logger.info(f"PDF extraction complete: {len(image_texts)} images and {len(table_texts)} tables processed")
            
        return combined_text, None
        
    except Exception as e:
        logger.exception(f"Error extracting text from PDF: {e}")
        return None, f"Error processing PDF file: {str(e)}"

# Helper function to extract text from DOCX files
def extract_text_from_docx(file_path):
    """
    Extract text from DOCX file including tables and images
    
    Uses python-docx for text and table extraction
    Extracts embedded images and processes them with OCR if possible
    """
    try:
        text = ""
        table_texts = []
        image_texts = []
        
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"DOCX file not found: {file_path}")
            return None, "DOCX file not found. Please upload the file again."
        
        # Check if python-docx is properly installed
        if isinstance(docx, DocxPlaceholder):
            logger.error("python-docx module is not installed")
            return None, "DOCX processing is currently unavailable. The server is missing required libraries."
            
        # Open and extract text from the DOCX
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract and format tables
        for table_idx, table in enumerate(doc.tables):
            table_text = ""
            
            # Process each row in the table
            for row in table.rows:
                row_texts = []
                
                # Process each cell in the row
                for cell in row.cells:
                    # Get all text from the cell (including from any paragraphs within)
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        cell_text += paragraph.text.strip() + " "
                    
                    row_texts.append(cell_text.strip())
                
                # Add formatted row to table text
                if any(row_texts):  # Only add non-empty rows
                    table_text += " | ".join(row_texts) + "\n"
            
            # Add table if it contains content
            if table_text.strip():
                table_texts.append({
                    'index': table_idx + 1,
                    'text': table_text.strip()
                })
                logger.info(f"Extracted table {table_idx+1} from DOCX")
        
        # Extract images from DOCX if possible
        try:
            import zipfile
            from xml.dom.minidom import parseString
            import os.path
            
            # DOCX files are ZIP archives
            image_index = 0
            with zipfile.ZipFile(file_path) as docx_zip:
                # Get list of all files in the docx
                file_list = docx_zip.namelist()
                
                # Find image files
                image_files = [f for f in file_list if f.startswith('word/media/')]
                
                for img_file in image_files:
                    try:
                        # Extract image to temp file
                        temp_img_path = os.path.join(
                            os.path.dirname(file_path),
                            f"temp_docx_img_{image_index}.jpg"
                        )
                        
                        with open(temp_img_path, "wb") as temp_file:
                            temp_file.write(docx_zip.read(img_file))
                        
                        # Check image size - skip tiny images (likely icons)
                        img = Image.open(temp_img_path)
                        width, height = img.size
                        
                        # Skip small images (likely icons or decorative elements)
                        if width < 100 or height < 100:
                            os.remove(temp_img_path)
                            continue
                        
                        # Process with OCR if pytesseract is available
                        if not isinstance(pytesseract, PytesseractPlaceholder):
                            # Try OCR on the image
                            ocr_text = pytesseract.image_to_string(img)
                            
                            # Only add if meaningful text was found (more than 10 chars)
                            if ocr_text and len(ocr_text.strip()) > 10:
                                image_texts.append({
                                    'index': image_index + 1,
                                    'text': ocr_text.strip()
                                })
                                logger.info(f"Extracted OCR text from image {image_index+1} in DOCX")
                                image_index += 1
                        
                        # Clean up the temporary file
                        os.remove(temp_img_path)
                        
                    except Exception as img_err:
                        logger.warning(f"Error processing DOCX image {img_file}: {str(img_err)}")
                        # Continue with next image
            
        except Exception as img_extract_err:
            logger.warning(f"Error extracting images from DOCX: {str(img_extract_err)}")
            # Continue with processing - image extraction is optional enhancement
        
        # Combine all extracted content
        combined_text = text
        
        # Add table texts
        for table_data in table_texts:
            combined_text += f"\n\n[Extracted from table {table_data['index']}]\n{table_data['text']}"
        
        # Add image texts
        for img_data in image_texts:
            combined_text += f"\n\n[Extracted from image {img_data['index']}]\n{img_data['text']}"
        
        # Check if we got any text
        if not combined_text or len(combined_text.strip()) == 0:
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return None, "No readable text found in the DOCX file."
        
        # Log extraction stats
        logger.info(f"DOCX extraction complete: {len(image_texts)} images with text and {len(table_texts)} tables processed")
            
        return combined_text, None
        
    except Exception as e:
        logger.exception(f"Error extracting text from DOCX: {e}")
        return None, f"Error processing DOCX file: {str(e)}"

# Helper function to generate summary using OpenAI
def generate_summary(text):
    """Generate a summary of the text using OpenAI API"""
    try:
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, "OpenAI API key not configured. Please contact the administrator."
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create prompt for summarization
        prompt = f"""Please summarize the following document in a detailed yet concise manner. 
        Focus on the main points, key arguments, and significant findings.
        Maintain the original context and meaning while condensing the content.
        
        DOCUMENT:
        {text[:10000]}  # Limiting input to first 10000 chars to avoid token limits
        
        SUMMARY:"""
        
        # Generate summary
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.3
        )
        
        summary = response.choices[0].message.content.strip()
        return summary, None
        
    except Exception as e:
        logger.exception(f"Error generating summary: {e}")
        return None, f"Error generating summary: {str(e)}"

# Generate personalized interaction tips
def generate_flashcards(document_text, document_id):
    """
    Generate educational flashcards based on document content
    
    Args:
        document_text: The full text content of the document
        document_id: The document ID for storing in session
        
    Returns:
        List of flashcards in JSON format, or None and error message if generation fails
    """
    try:
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, "OpenAI API key not configured. Please contact the administrator."
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create prompt for flashcard generation
        prompt = f"""Generate 5-10 educational flashcards based on the following document content.
        Each flashcard should have a question on one side and the answer on the other.
        Focus on key concepts, definitions, and important facts from the document.
        Make the flashcards educational and helpful for someone studying this material.
        
        Return the flashcards in the following JSON format:
        [
          {{"question": "Question 1?", "answer": "Answer 1"}},
          {{"question": "Question 2?", "answer": "Answer 2"}},
          ...
        ]
        
        DOCUMENT CONTENT:
        {document_text[:4000]}  # Limiting input to first 4000 chars to avoid token limits
        """
        
        # Generate flashcards
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=2000,
            temperature=0.7
        )
        
        result = response.choices[0].message.content.strip()
        
        # Parse the JSON response
        try:
            flashcards_data = json.loads(result)
            
            # Make sure we have the expected structure
            if 'flashcards' in flashcards_data:
                flashcards = flashcards_data['flashcards']
            else:
                # Try to find an array directly in the response
                for key, value in flashcards_data.items():
                    if isinstance(value, list) and len(value) > 0:
                        flashcards = value
                        break
                else:
                    # If no array found within objects, check if the response itself is an array
                    if isinstance(flashcards_data, list):
                        flashcards = flashcards_data
                    else:
                        flashcards = []
            
            # Validate each flashcard has question and answer
            valid_flashcards = []
            for card in flashcards:
                if isinstance(card, dict) and 'question' in card and 'answer' in card:
                    valid_flashcards.append({
                        'question': card['question'],
                        'answer': card['answer']
                    })
            
            if not valid_flashcards:
                return None, "Failed to generate valid flashcards from the document content."
                
            # Initialize flashcards session storage if needed
            if 'flashcards' not in session:
                session['flashcards'] = {}
                
            # Store flashcards in session keyed by document ID
            session['flashcards'][str(document_id)] = valid_flashcards
            session.modified = True
                
            return valid_flashcards, None
            
        except json.JSONDecodeError:
            logger.error(f"Failed to parse flashcards JSON: {result}")
            return None, "Failed to generate valid flashcards. The AI returned malformed data."
            
    except Exception as e:
        logger.exception(f"Error generating flashcards: {e}")
        return None, f"Error generating flashcards: {str(e)}"

def generate_quiz(document_text, document_id):
    """
    Generate a multiple-choice quiz based on document content
    
    Args:
        document_text: The full text content of the document
        document_id: The document ID for storing in session
        
    Returns:
        List of quiz questions in JSON format, or None and error message if generation fails
    """
    try:
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, "OpenAI API key not configured. Please contact the administrator."
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create an improved system prompt with stricter formatting requirements
        system_prompt = """You are an education quiz generator that creates high-quality multiple choice questions based on documents. 
        
        IMPORTANT: Generate exactly 5 multiple choice questions based on the document provided. Each question should:
        1. Be clearly based on information found in the document
        2. Have exactly 4 answer options labeled A, B, C, and D
        3. Have exactly one correct answer
        4. Focus on important concepts rather than trivial details
        
        You MUST format your response as a valid JSON array of objects with these exact keys:
        - "question": the text of the question (string)
        - "options": an object with keys A, B, C, D containing the text of each option (object)
        - "correct_answer": the letter of the correct option (string, one of: "A", "B", "C", or "D")
        
        Required format example:
        [
            {
                "question": "What is the capital of France?",
                "options": {
                    "A": "London",
                    "B": "Berlin",
                    "C": "Paris",
                    "D": "Madrid"
                },
                "correct_answer": "C"
            }
        ]
        
        DO NOT include any explanation or additional text outside the JSON array.
        DO NOT use any other format.
        """
        
        # Create user prompt with truncated document text
        doc_text_truncated = document_text[:4000]  # Limiting input to first 4000 chars to avoid token limits
        user_prompt = f"""Please generate a multiple-choice quiz based on this document content.
        Create questions that test understanding of key concepts, important facts, and main ideas.
        The quiz should help students review and reinforce their understanding of the material.
        
        Document content:
        {doc_text_truncated}
        
        Remember to follow the format specified in the system instructions exactly.
        """
        
        # Generate quiz
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            max_tokens=2000,
            temperature=0.7
        )
        
        result = response.choices[0].message.content.strip()
        
        # Parse the JSON response
        try:
            quiz_data = json.loads(result)
            
            # Make sure we have the expected structure
            quiz_questions = []
            if isinstance(quiz_data, list):
                quiz_questions = quiz_data
            elif any(key in quiz_data for key in ['quiz', 'questions']):
                for key in ['quiz', 'questions']:
                    if key in quiz_data and isinstance(quiz_data[key], list):
                        quiz_questions = quiz_data[key]
                        break
            else:
                # Try to find an array directly in the response
                for key, value in quiz_data.items():
                    if isinstance(value, list) and len(value) > 0:
                        quiz_questions = value
                        break
            
            # Validate each quiz question with improved validation logic
            valid_questions = []
            for question in quiz_questions:
                # Skip if not a dictionary
                if not isinstance(question, dict):
                    logger.warning(f"Skipping quiz question that is not a dict: {type(question)}")
                    continue
                
                # Check for required fields
                if not all(key in question for key in ['question', 'options', 'correct_answer']):
                    logger.warning(f"Skipping quiz question missing required fields: {question.keys()}")
                    continue
                
                # If options isn't a dict, try to convert it
                if not isinstance(question['options'], dict):
                    logger.warning(f"Question options is not a dict, attempting to fix: {type(question['options'])}")
                    # Try to convert options to a dict if it's a list
                    if isinstance(question['options'], list) and len(question['options']) >= 4:
                        options_dict = {
                            'A': question['options'][0],
                            'B': question['options'][1],
                            'C': question['options'][2],
                            'D': question['options'][3]
                        }
                        question['options'] = options_dict
                    else:
                        # Can't convert, skip this question
                        continue
                
                # Check if options has at least the keys A, B, C, D
                if not all(option in question['options'] for option in ['A', 'B', 'C', 'D']):
                    logger.warning(f"Skipping quiz question with invalid options keys: {question['options'].keys()}")
                    continue
                
                # Check if correct_answer is valid
                if question['correct_answer'] not in ['A', 'B', 'C', 'D']:
                    logger.warning(f"Skipping quiz question with invalid correct_answer: {question['correct_answer']}")
                    continue
                
                # If we got here, the question is valid
                valid_questions.append({
                    'question': question['question'],
                    'options': question['options'],
                    'correct_answer': question['correct_answer']
                })
            
            if not valid_questions:
                return None, "Failed to generate valid quiz questions from the document content."
                
            # Initialize quizzes session storage if needed
            if 'quizzes' not in session:
                session['quizzes'] = {}
                
            # Store quiz in session keyed by document ID
            session['quizzes'][str(document_id)] = valid_questions
            session.modified = True
            
            logger.info(f"Successfully generated quiz with {len(valid_questions)} questions for document {document_id}")
            return valid_questions, None
            
        except json.JSONDecodeError:
            logger.error(f"Failed to parse quiz JSON: {result}")
            return None, "Failed to generate valid quiz questions. The AI returned malformed data."
            
    except Exception as e:
        logger.exception(f"Error generating quiz: {e}")
        return None, f"Error generating quiz: {str(e)}"

def generate_interaction_tips(document_summary, filetype):
    """
    Generate personalized document interaction tips based on the document summary
    
    Args:
        document_summary: The summary of the document
        filetype: The type of document (PDF, DOCX, etc.)
        
    Returns:
        List of interaction tips, or None and error message if generation fails
    """
    try:
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, "OpenAI API key not configured. Please contact the administrator."
            
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create system prompt for generating tips
        system_prompt = """You are an expert educational guide that helps users interact with documents effectively.
        Your goal is to provide personalized, insightful tips that help the user extract maximum value from their document through AI interactions.
        Focus on making tips that are:
        1. Specific to the document's content and subject matter
        2. Written in a friendly, encouraging tone
        3. Actionable and practical
        4. Educational and promote deeper learning
        """
        
        # Create user prompt for generating tips
        user_prompt = f"""Based on the following document summary, create 5 personalized tips to help the user interact with and learn from this document.
        These tips should help the user ask better questions, explore key concepts, and get maximum value from the AI chat feature.
        
        For each tip, provide:
        1. A short, catchy title (3-5 words) that encapsulates the tip
        2. An emoji that represents the tip (choose from 📚 🔍 💡 🗣️ 🧠 📝 🔄 📊 🔬 🧩)
        3. A concise explanation (1-2 sentences) of how to apply this tip
        
        Format your response as a list of 5 numbered tips with the following structure for each tip:
        
        **[NUMBER]. [EMOJI] [TITLE]**
        [EXPLANATION]
        
        Document type: {filetype}
        Document summary: {document_summary[:1000]}
        """
        
        # Generate tips
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        tips = response.choices[0].message.content.strip()
        
        # Log the tips for debugging
        logger.info(f"Generated interaction tips: {tips[:100]}...")
        
        return tips, None
        
    except Exception as e:
        logger.exception(f"Error generating interaction tips: {e}")
        return None, f"Error generating tips: {str(e)}"

def generate_learning_suggestions(document_summary):
    """
    Generate next learning suggestions based on the document content
    
    Args:
        document_summary: The summary of the document
        
    Returns:
        List of learning suggestions as a JSON string, or None and error message if generation fails
    """
    try:
        # Get OpenAI API key from environment variables
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, "OpenAI API key not configured. Please contact the administrator."
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create prompt for generating learning suggestions
        prompt = f"""Based on this document summary, suggest 3 to 4 related educational topics the user should study next.
        Return them as a JSON array of objects with 'title' and 'description' properties.
        
        Make sure the suggestions are specific and related to the document content, not generic learning topics.
        Each suggestion should include a concise title (3-5 words) and a brief description (15-25 words).
        
        Document Summary:
        {document_summary[:2000]}  # Limiting to first 2000 chars
        
        Example format:
        [
            {{
                "title": "Advanced Machine Learning",
                "description": "Explore mathematical foundations and implementations of neural networks, decision trees, and support vector machines."
            }},
            {{
                "title": "Data Visualization Techniques",
                "description": "Learn effective methods for representing complex datasets through charts, graphs, and interactive visualizations."
            }}
        ]
        
        Response:"""
        
        # Generate suggestions
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=600,
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        
        suggestions_json = response.choices[0].message.content.strip()
        # Parse to validate JSON and then return the string for storage
        json.loads(suggestions_json)  # Validate JSON
        return suggestions_json, None
        
    except json.JSONDecodeError as e:
        logger.exception(f"Invalid JSON in learning suggestions: {e}")
        return None, f"Failed to generate valid learning suggestions format: {str(e)}"
    except Exception as e:
        logger.exception(f"Error generating learning suggestions: {e}")
        return None, f"Failed to generate learning suggestions: {str(e)}"
        
def translate_document(document_text, target_language):
    """
    Translate document text to target language and provide language character guides
    
    Args:
        document_text: Text to translate
        target_language: Target language code (e.g., 'es', 'fr', 'de', 'zh')
        
    Returns:
        Tuple of (translated_text, language_guide, error_message)
    """
    try:
        # Get OpenAI API key from environment variables
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, None, "OpenAI API key not configured. Please contact the administrator."
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Define language information
        language_info = {
            'es': {'name': 'Spanish', 'native': 'Español'},
            'fr': {'name': 'French', 'native': 'Français'},
            'de': {'name': 'German', 'native': 'Deutsch'},
            'it': {'name': 'Italian', 'native': 'Italiano'},
            'pt': {'name': 'Portuguese', 'native': 'Português'},
            'ru': {'name': 'Russian', 'native': 'Русский'},
            'zh': {'name': 'Chinese (Simplified)', 'native': '中文'},
            'ja': {'name': 'Japanese', 'native': '日本語'},
            'ko': {'name': 'Korean', 'native': '한국어'},
            'ar': {'name': 'Arabic', 'native': 'العربية'},
            'hi': {'name': 'Hindi', 'native': 'हिन्दी'},
            'en': {'name': 'English', 'native': 'English'},
        }
        
        target_lang_name = language_info.get(target_language, {}).get('name', target_language)
        target_lang_native = language_info.get(target_language, {}).get('native', target_language)
        
        # Create prompt for translation
        trans_prompt = f"""Translate the following text to {target_lang_name} ({target_lang_native}).
        Maintain the original formatting as much as possible.
        
        Text to translate:
        {document_text[:4000]}  # Limiting to first 4000 chars for token constraints
        
        Translated text:"""
        
        # Generate translation
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        trans_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional translator."},
                {"role": "user", "content": trans_prompt}
            ],
            max_tokens=4000,
            temperature=0.3
        )
        
        translated_text = trans_response.choices[0].message.content.strip()
        
        # Create prompt for language guide
        guide_prompt = f"""Create a concise language guide for {target_lang_name} with the following sections:
        
        1. Alphabet/Character information (especially if non-Latin)
        2. Basic pronunciation rules (3-4 key points)
        3. Common phrases (5-6 everyday phrases with pronunciation guides)
        
        Format the guide in a structured way with clear headings. Keep it brief but useful for someone encountering this language for the first time."""
        
        # Generate language guide
        guide_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a language education expert."},
                {"role": "user", "content": guide_prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        language_guide = guide_response.choices[0].message.content.strip()
        
        return translated_text, language_guide, None
        
    except Exception as e:
        logger.exception(f"Error translating document: {e}")
        return None, None, f"Failed to translate document: {str(e)}"

# Simple text similarity search function

def simple_text_search(query, document_text, chunk_size=1000, overlap=200, top_k=3):
    """
    Perform a simple text similarity search to find relevant document chunks
    
    Args:
        query: The search query (user question)
        document_text: The document text to search within
        chunk_size: Size of each document chunk
        overlap: Overlap between chunks
        top_k: Number of most relevant chunks to return
    
    Returns:
        List of most relevant text chunks
    """
    try:
        # Normalize query (lowercase, remove extra spaces)
        query = " ".join(query.lower().split())
        
        # Break document into overlapping chunks
        chunks = []
        doc_length = len(document_text)
        
        for i in range(0, doc_length, chunk_size - overlap):
            chunk = document_text[i:i + chunk_size]
            if len(chunk) < 100:  # Skip very small chunks
                continue
            chunks.append(chunk)
        
        # If we have no chunks, return the whole document
        if not chunks:
            return [document_text[:5000]]  # Return first 5000 chars
            
        # Calculate simple relevance score for each chunk
        chunk_scores = []
        query_terms = set(query.split())
        
        for i, chunk in enumerate(chunks):
            chunk_lower = chunk.lower()
            
            # Count how many query terms appear in the chunk
            term_matches = sum(1 for term in query_terms if term in chunk_lower)
            
            # Count exact phrase matches (more weight)
            phrase_matches = chunk_lower.count(query) * 3
            
            # Calculate final score
            score = term_matches + phrase_matches
            
            chunk_scores.append((i, score, chunk))
        
        # Sort by score (highest first) and take top_k
        chunk_scores.sort(key=lambda x: x[1], reverse=True)
        top_chunks = [chunk for _, _, chunk in chunk_scores[:top_k]]
        
        # If no chunks matched, return first chunk as fallback
        if not top_chunks or all(score == 0 for _, score, _ in chunk_scores[:top_k]):
            return [chunks[0]]
            
        return top_chunks
        
    except Exception as e:
        logger.exception(f"Error in similarity search: {e}")
        # Fallback to first 5000 chars if search fails
        return [document_text[:5000]]

# Generate response for chat using OpenAI
def generate_chat_response(document_id, user_message):
    """Generate a response to a user question about a document using OpenAI API"""
    try:
        # Initialize document text
        document_text = None
        
        # Check if this is a session-only document
        is_session_doc = str(document_id).startswith('temp_')
        
        # Logic to retrieve document text (from DB or session)
        if is_session_doc:
            # Look for document in session
            for doc in session.get('uploads', []):
                if str(doc.get('id')) == str(document_id):
                    document_text = doc.get('text_content')
                    logger.info(f"Using session text for chat with document: {document_id}")
                    break
        else:
            # Try to retrieve from database
            try:
                # Convert to int for database query
                document_id_int = int(document_id)
                document = Document.query.get(document_id_int)
                if document and document.text_content:
                    document_text = document.text_content
                else:
                    # Fall back to session if DB document not found
                    for doc in session.get('uploads', []):
                        if str(doc.get('id')) == str(document_id):
                            document_text = doc.get('text_content')
                            logger.info(f"Using session text as fallback for chat: {document_id}")
                            break
            except (ValueError, Exception) as db_error:
                logger.error(f"Error retrieving document from database: {db_error}")
                # Try session as fallback
                for doc in session.get('uploads', []):
                    if str(doc.get('id')) == str(document_id):
                        document_text = doc.get('text_content')
                        break
        
        # Check if we have document text
        if not document_text:
            return None, "Document not found or has no content"
        
        # Check for API key
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            return None, "OpenAI API key not configured. Please contact the administrator."
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Perform text similarity search to find relevant chunks
        relevant_chunks = simple_text_search(
            user_message, 
            document_text,
            chunk_size=1500,
            overlap=300, 
            top_k=3
        )
        
        # Combine relevant chunks for context (limit total to ~5000 chars)
        combined_chunks = " ".join(relevant_chunks)
        if len(combined_chunks) > 5000:
            combined_chunks = combined_chunks[:5000]
        
        # Create system message with enhanced instructions and context
        system_message = f"""You are an AI assistant answering questions about a specific document.
        The following are the most relevant sections of the document related to the user's question.
        
        DOCUMENT CONTENT:
        {combined_chunks}
        
        Instructions:
        1. Only use the document content provided above to answer questions.
        2. If the answer cannot be found in these document sections, say so politely.
        3. Do not make up information or use external knowledge.
        4. If the document content is in a language other than English, answer in that same language.
        5. Format your response neatly with paragraphs where appropriate.
        """
        
        # Generate response
        # the newest OpenAI model is "gpt-4o" which was released May 13, 2024.
        # do not change this unless explicitly requested by the user
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message}
            ],
            max_tokens=800,
            temperature=0.3
        )
        
        answer = response.choices[0].message.content.strip()
        return answer, None
        
    except Exception as e:
        logger.exception(f"Error generating chat response: {e}")
        return None, f"Error generating response: {str(e)}"

# Routes
@app.route('/')
def index():
    """Home page with file upload form"""
    # Generate a session ID if none exists
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        
    # Initialize session storage for uploads and chat history if not exists
    if 'uploads' not in session:
        session['uploads'] = []
        
    if 'chat_history' not in session:
        session['chat_history'] = {}
        
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
@db_retry(max_retries=3)
def upload_file():
    """Handle file upload"""
    # Check if a file was uploaded
    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('index'))
        
    file = request.files['file']
    
    # Check if a file was selected
    if file.filename == '':
        flash('No file selected')
        return redirect(url_for('index'))
        
    # Check file type
    if not allowed_file(file.filename):
        flash('Invalid file type. Please upload a PDF, DOCX, TXT, JPG, JPEG, or PNG file')
        return redirect(url_for('index'))
    
    try:
        # Create upload folder if it doesn't exist
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Save the file with a unique filename to prevent overwrites
        filename = secure_filename(file.filename)
        unique_id = str(uuid.uuid4())[:8]
        base_name, extension = os.path.splitext(filename)
        unique_filename = f"{base_name}_{unique_id}{extension}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        
        # Extract text based on file type
        filetype = extension[1:].lower()  # Remove the dot
        if filetype == 'pdf':
            extracted_text, error = extract_text_from_pdf(file_path)
        elif filetype == 'docx':
            extracted_text, error = extract_text_from_docx(file_path)
        elif filetype == 'txt':
            extracted_text, error = extract_text_from_txt(file_path)
        elif filetype in ['jpg', 'jpeg', 'png']:
            # Inform user about OCR processing through the logger
            logger.info(f"Processing image with OCR: {filename}")
            
            # Extract text using OCR
            extracted_text, error = extract_text_from_image(file_path)
            
            # Add OCR indicator to inform the user
            if extracted_text and not error:
                ocr_info = (
                    "[TEXT EXTRACTED USING OCR TECHNOLOGY]\n\n"
                    "This text was extracted from your image using Optical Character Recognition (OCR) "
                    "with multi-language support. The quality of extraction depends on image clarity "
                    "and text formatting. You can ask questions about this content just like with any "
                    "text document.\n\n"
                    "---------------------------------------------------\n\n"
                )
                extracted_text = ocr_info + extracted_text
                logger.info(f"Successfully extracted text from image: {filename}")
        else:
            # This shouldn't happen due to allowed_file check, but just in case
            extracted_text, error = None, "Unsupported file type"
        
        # Handle extraction errors
        if error:
            flash(f"Error: {error}")
            return redirect(url_for('index'))
        
        # Generate summary
        summary, summary_error = generate_summary(extracted_text)
        
        # Handle summary generation errors
        if summary_error:
            logger.warning(f"Summary generation error: {summary_error}")
            summary = "Summary generation failed. You can still view and chat with the document content."
            flash("Warning: Could not generate summary due to an error, but the document was processed.")
            
        # Generate interaction tips
        tips, tips_error = generate_interaction_tips(summary, filetype)
        
        # If tips generation failed, log the error but continue (non-critical)
        if tips_error:
            logger.warning(f"Error generating interaction tips: {tips_error}")
            tips = '["Ask about key topics", "Request explanations of important concepts"]'
            
        # Generate learning suggestions
        suggestions, suggestions_error = generate_learning_suggestions(summary)
        
        # If suggestions generation failed, log the error but continue (non-critical)
        if suggestions_error:
            logger.warning(f"Error generating learning suggestions: {suggestions_error}")
            # Default suggestions as JSON string
            suggestions = '[{"title":"Related Topic 1","description":"Explore similar concepts related to this document content."},{"title":"Advanced Learning","description":"Study more advanced aspects of this subject area."}]'
        
        # Initialize document variable
        document = None
        document_id = None
        
        # Store in database with improved error handling
        try:
            # Make sure session ID exists
            if 'session_id' not in session:
                session['session_id'] = str(uuid.uuid4())
                
            document = Document(
                session_id=session['session_id'],
                filename=filename,  # Use original filename for display
                filetype=filetype.upper(),
                summary=summary,
                text_content=extracted_text,
                interaction_tips=tips,
                learning_suggestions=suggestions,
                upload_time=datetime.now()
            )
            db.session.add(document)
            db.session.commit()
            document_id = document.id
            logger.info(f"Document saved to database: {filename} (ID: {document_id})")
            
        except Exception as db_error:
            logger.error(f"Database error when saving document: {db_error}")
            db.session.rollback()
            
            # Generate a temporary document ID for session storage
            document_id = f"temp_{uuid.uuid4()}"
            
            flash("Warning: Document saved to session only due to database connection issues. Some features may be limited.", "warning")
        
        # Initialize session storage if needed
        if 'uploads' not in session:
            session['uploads'] = []
            
        if 'chat_history' not in session:
            session['chat_history'] = {}
        
        # Store document ID in session
        session['last_document_id'] = document_id
        
        # Store document metadata and content in session for fallback access
        upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        session_document = {
            'id': document_id,
            'filename': filename,
            'filetype': filetype.upper(),
            'summary': summary[:200] + "..." if summary and len(summary) > 200 else summary,
            'upload_time': upload_time,
            # Store more data for fallback use
            'full_summary': summary,
            'text_content': extracted_text,
            'interaction_tips': tips,
            'learning_suggestions': suggestions,
            'storage_path': file_path
        }
        
        # Add to the beginning of the list (most recent first)
        session['uploads'].insert(0, session_document)
        
        # Initialize empty chat history for this document
        session['chat_history'][str(document_id)] = []
        
        # Make sure to persist the session
        session.modified = True
        
        # Redirect to document view
        return redirect(url_for('view_document', document_id=document_id))
        
    except Exception as e:
        logger.exception(f"Error processing file upload: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        flash(f"Error: {str(e)}")
        return redirect(url_for('index'))

@app.route('/document/<document_id>')
@db_retry(max_retries=3)
def view_document(document_id):
    """View uploaded document with summary and chat"""
    try:
        # Check if this might be a temporary document ID from session storage
        if str(document_id).startswith('temp_'):
            # This is a session-only document, so we need to get it from session
            document_found = False
            session_document = None
            
            # Check if we have session uploads
            if 'uploads' in session:
                for doc in session['uploads']:
                    if str(doc['id']) == document_id:
                        document_found = True
                        session_document = doc
                        break
            
            if not document_found:
                flash("Document not found in session storage")
                return redirect(url_for('dashboard'))
            
            # Create a document-like object from session data for the template
            class SessionDocument:
                def __init__(self, doc_dict):
                    self.id = doc_dict['id']
                    self.filename = doc_dict['filename']
                    self.filetype = doc_dict['filetype']
                    self.summary = doc_dict.get('full_summary', doc_dict.get('summary', 'No summary available'))
                    self.text_content = doc_dict.get('text_content', '')
                    self.interaction_tips = doc_dict.get('interaction_tips', '[]')
                    self.learning_suggestions = doc_dict.get('learning_suggestions', '[]')
                    self.upload_time = doc_dict.get('upload_time', 'Unknown')
                    self.session_id = session.get('session_id', '')
            
            document = SessionDocument(session_document)
            logger.info(f"Viewing session-only document: {document_id}")
            
            # Initialize chat history in session if not exists
            if 'chat_history' not in session:
                session['chat_history'] = {}
            
            # If we don't have chat history for this document in session, initialize it
            doc_id_str = str(document_id)
            if doc_id_str not in session['chat_history']:
                session['chat_history'][doc_id_str] = []
                session.modified = True
            
            # Check if there are flashcards for this document
            has_flashcards = False
            if 'flashcards' in session and doc_id_str in session['flashcards']:
                has_flashcards = True
            
            # For session-only documents, there won't be DB chat messages
            return render_template(
                'document.html',
                document=document,
                chat_messages=[],
                session_chat_history=session['chat_history'].get(doc_id_str, []),
                has_flashcards=has_flashcards,
                is_session_only=True,
                translation_languages=TRANSLATION_LANGUAGES,
                view_translation=False,
                language_code=None,
                language_info={}
            )
        else:
            # Regular DB document - convert to int for database query
            try:
                document_id = int(document_id)
            except ValueError:
                flash("Invalid document ID")
                return redirect(url_for('dashboard'))
                
            # Retrieve document with error handling
            document = Document.query.get_or_404(document_id)
                
            # Check if user owns this document (session-based)
            if document.session_id != session.get('session_id'):
                flash("You don't have permission to view this document")
                return redirect(url_for('index'))
            
            # Retrieve chat history from database with error handling
            chat_messages = []
            try:
                chat_messages = ChatMessage.query.filter_by(
                    document_id=document_id,
                    session_id=session['session_id']
                ).order_by(ChatMessage.created_at).all()
                logger.info(f"Retrieved {len(chat_messages)} chat messages for document {document_id}")
            except Exception as chat_error:
                logger.error(f"Error retrieving chat messages: {chat_error}")
                db.session.rollback()
                flash("Unable to retrieve chat history. Using cached data if available.", "warning")
            
            # Initialize chat history in session if not exists
            if 'chat_history' not in session:
                session['chat_history'] = {}
            
            # If we don't have chat history for this document in session, initialize it
            doc_id_str = str(document_id)
            if doc_id_str not in session['chat_history']:
                session['chat_history'][doc_id_str] = []
                session.modified = True
            
            # Convert database messages to dict format for template rendering
            db_messages = []
            try:
                db_messages = [msg.to_dict() for msg in chat_messages]
            except Exception as convert_error:
                logger.error(f"Error converting chat messages: {convert_error}")
            
            # Check if there are flashcards for this document
            has_flashcards = False
            if 'flashcards' in session and doc_id_str in session['flashcards']:
                has_flashcards = True
            
            return render_template(
                'document.html',
                document=document,
                chat_messages=db_messages,
                session_chat_history=session['chat_history'].get(doc_id_str, []),
                has_flashcards=has_flashcards,
                translation_languages=TRANSLATION_LANGUAGES,
                view_translation=False,
                language_code=None,
                language_info={}
            )
        
    except Exception as e:
        # Log the error
        logger.exception(f"Error viewing document {document_id}: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Check if we can find this document in the session as a fallback
        if 'uploads' in session:
            for doc in session['uploads']:
                if str(doc['id']) == str(document_id):
                    flash(f"Using cached document data due to database error: {str(e)}", "warning")
                    logger.info(f"Falling back to session data for document: {document_id}")
                    
                    # Redirect to the same route which will now handle it as a session document 
                    # if it matches the session-only format
                    return redirect(url_for('view_document', document_id=doc['id']))
        
        # Show error message and redirect to dashboard
        flash(f"An error occurred while viewing the document: {str(e)}", "error")
        return redirect(url_for('dashboard'))

@app.route('/api/chat/<document_id>', methods=['POST'])
@db_retry(max_retries=3)
def chat_with_document(document_id):
    """API endpoint for chat functionality"""
    try:
        # Check if user is in session
        if 'session_id' not in session:
            return jsonify({'error': 'Session expired. Please refresh the page.'}), 401
        
        # Check if it's a JSON request or a file upload (multipart/form-data)
        if request.is_json:
            # Handle standard text message
            data = request.get_json()
            if not data or 'message' not in data:
                return jsonify({'error': 'No message provided'}), 400
            
            user_message = data['message'].strip()
            if not user_message:
                return jsonify({'error': 'Empty message'}), 400
                
            has_attachment = False
            attachment_filename = None
            attachment_original_filename = None
            attachment_type = None
            attachment_text = None
        else:
            # This might be a file upload with form data
            user_message = request.form.get('message', '').strip()
            if not user_message and 'file' not in request.files:
                return jsonify({'error': 'No message or file provided'}), 400
                
            # Initialize attachment variables
            has_attachment = False
            attachment_filename = None
            attachment_original_filename = None
            attachment_type = None
            attachment_text = None
            
            # Check if there's a file attached
            if 'file' in request.files:
                file = request.files['file']
                if file and file.filename and allowed_file(file.filename):
                    # Process the attachment
                    has_attachment = True
                    original_filename = secure_filename(file.filename)
                    filename_parts = os.path.splitext(original_filename)
                    extension = filename_parts[1].lower()
                    
                    # Generate a unique filename
                    unique_filename = f"{filename_parts[0]}_{uuid.uuid4().hex}{extension}"
                    file_path = os.path.join(app.config['ATTACHMENT_FOLDER'], unique_filename)
                    
                    # Save the file
                    file.save(file_path)
                    
                    # Set attachment information
                    attachment_filename = unique_filename
                    attachment_original_filename = original_filename
                    
                    # Determine attachment type and extract text if possible
                    if extension in ['.pdf']:
                        attachment_type = 'pdf'
                        extracted_text, error = extract_text_from_pdf(file_path)
                        if extracted_text:
                            attachment_text = extracted_text
                    elif extension in ['.docx']:
                        attachment_type = 'docx'
                        extracted_text, error = extract_text_from_docx(file_path)
                        if extracted_text:
                            attachment_text = extracted_text
                    elif extension in ['.txt']:
                        attachment_type = 'txt'
                        extracted_text, error = extract_text_from_txt(file_path)
                        if extracted_text:
                            attachment_text = extracted_text
                    elif extension in ['.jpg', '.jpeg', '.png']:
                        attachment_type = 'image'
                        # Log OCR processing for attachment
                        logger.info(f"Processing attachment image with OCR: {original_filename}")
                        extracted_text, error = extract_text_from_image(file_path)
                        if extracted_text:
                            # Add OCR header for image attachments
                            ocr_note = "[Text extracted from the attached image using OCR technology]\n\n"
                            attachment_text = ocr_note + extracted_text
                            logger.info(f"Successfully extracted text from image attachment: {original_filename}")
        
        # Check if this is a session-only document
        is_session_doc = str(document_id).startswith('temp_')
        document = None
        document_text = None
        
        if is_session_doc:
            # This is a session-only document, so get it from session
            session_document = None
            for doc in session.get('uploads', []):
                if str(doc['id']) == str(document_id):
                    session_document = doc
                    break
                    
            if not session_document:
                return jsonify({'error': 'Document not found in session'}), 404
            
            # For session documents, use the stored text content
            document_text = session_document.get('text_content')
            if not document_text:
                return jsonify({
                    'error': 'This document has no extractable text content',
                    'details': 'The file may be an image-based PDF or contain unsupported formatting.'
                }), 422
            
            # Create a mock document object with the document_id attribute for later use
            class SessionDocument:
                def __init__(self, doc_id, text):
                    self.id = doc_id
                    self.text_content = text
                    self.session_id = session.get('session_id', '')
                
                def to_dict(self):
                    return {
                        'id': self.id,
                        'text_content': self.text_content,
                        'session_id': self.session_id
                    }
            
            document = SessionDocument(document_id, document_text)
            logger.info(f"Using session document for chat: {document_id}")
        else:
            # Try to convert to integer for database lookup
            try:
                document_id_int = int(document_id)
            except ValueError:
                return jsonify({'error': 'Invalid document ID'}), 400
            
            # Retrieve from database
            document = Document.query.get(document_id_int)
            if not document:
                # Check if we have it in session as a fallback
                for doc in session.get('uploads', []):
                    if str(doc['id']) == str(document_id):
                        document_text = doc.get('text_content')
                        if document_text:
                            document = SessionDocument(document_id, document_text)
                            logger.info(f"Using session document as fallback for chat: {document_id}")
                            is_session_doc = True
                            break
                
                if not document:
                    return jsonify({'error': 'Document not found'}), 404
            
            if not is_session_doc:
                # Check if document has content
                if not document.text_content:
                    return jsonify({
                        'error': 'This document has no extractable text content',
                        'details': 'The file may be an image-based PDF or contain unsupported formatting.'
                    }), 422
                
                # Check if user owns this document
                if document.session_id != session['session_id']:
                    return jsonify({'error': 'Unauthorized access'}), 403
        
        # Check for OpenAI API key
        if not os.environ.get("OPENAI_API_KEY"):
            return jsonify({
                'error': 'OpenAI API key not configured',
                'details': 'Please contact the administrator to set up the API key.'
            }), 503
        
        # Begin a database transaction to handle potential errors
        db.session.begin_nested()
        
        # Save user message with attachment info if present
        user_chat = ChatMessage(
            document_id=document_id,
            session_id=session['session_id'],
            message_type='user',
            content=user_message,
            has_attachment=has_attachment,
            attachment_filename=attachment_filename,
            attachment_original_filename=attachment_original_filename,
            attachment_type=attachment_type,
            attachment_text=attachment_text
        )
        db.session.add(user_chat)
        db.session.commit()
        
        # Prepare context for AI response
        # If there's attachment text, include it in the context
        additional_context = ""
        if attachment_text:
            additional_context = f"\n\nThe user has also attached a file with the following content:\n{attachment_text[:2000]}"
        
        # Generate AI response
        ai_response, error = generate_chat_response(document_id, user_message + additional_context)
        
        # Handle errors during response generation
        if error:
            # Create an error response message in chat
            error_message = f"I'm sorry, I couldn't process your request: {error}"
            
            ai_chat = ChatMessage(
                document_id=document_id,
                session_id=session['session_id'],
                message_type='assistant',
                content=error_message
            )
            db.session.add(ai_chat)
            db.session.commit()
            
            return jsonify({
                'user_message': user_chat.to_dict(),
                'ai_response': ai_chat.to_dict(),
                'warning': error
            }), 200  # Still return 200 to display the error in the chat
        
        # Save AI response
        ai_chat = ChatMessage(
            document_id=document_id,
            session_id=session['session_id'],
            message_type='assistant',
            content=ai_response
        )
        db.session.add(ai_chat)
        db.session.commit()
        
        # Store messages in session for history tracking
        doc_id_str = str(document_id)
        
        # Add user message to session history
        session['chat_history'][doc_id_str].append({
            'role': 'user',
            'content': user_message,
            'has_attachment': has_attachment,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
        
        # Add AI response to session history
        session['chat_history'][doc_id_str].append({
            'role': 'assistant',
            'content': ai_response,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
        
        # Make sure to persist the session
        session.modified = True
        
        # Return both messages
        return jsonify({
            'user_message': user_chat.to_dict(),
            'ai_response': ai_chat.to_dict()
        })
        
    except Exception as e:
        # Log the error
        logger.exception(f"Error in chat API: {e}")
        
        # Rollback any transaction in progress
        db.session.rollback()
        
        # Return error response
        return jsonify({
            'error': 'An unexpected error occurred',
            'details': str(e)
        }), 500

@app.route('/download/summary/<int:document_id>')
@db_retry(max_retries=3)
def download_summary(document_id):
    """Download document summary as text file"""
    try:
        # Retrieve document
        document = Document.query.get_or_404(document_id)
        
        # Check if user owns this document
        if document.session_id != session.get('session_id'):
            flash("You don't have permission to download this summary")
            return redirect(url_for('index'))
        
        # Create summary directory if it doesn't exist
        os.makedirs('static/summaries', exist_ok=True)
        
        # Create summary file
        summary_filename = f"summary_{document_id}.txt"
        summary_path = os.path.join('static/summaries', summary_filename)
        
        # Write summary to file
        with open(summary_path, 'w', encoding='utf-8') as f:
            f.write(f"Summary of: {document.filename}\n")
            f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(document.summary or "No summary available.")
        
        # Send file to user
        return send_from_directory(
            'static/summaries',
            summary_filename,
            as_attachment=True,
            download_name=f"Summary-{document.filename}.txt"
        )
    except Exception as e:
        # Log the error
        logger.exception(f"Error downloading summary: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Show error message
        flash(f"An error occurred while downloading the summary: {str(e)}")
        return redirect(url_for('document', document_id=document_id))

@app.route('/dashboard')
@db_retry(max_retries=3)
def dashboard():
    """User dashboard showing uploaded documents"""
    # Check if user is in session
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        
    # Initialize session storage if not exists
    if 'uploads' not in session:
        session['uploads'] = []
        
    if 'chat_history' not in session:
        session['chat_history'] = {}
        
    if 'flashcards' not in session:
        session['flashcards'] = {}
    
    # Retrieve user's documents from database with error handling
    db_documents = []
    try:
        db_documents = Document.query.filter_by(session_id=session['session_id']).order_by(Document.upload_time.desc()).all()
        logger.info(f"Retrieved {len(db_documents)} documents from database")
    except Exception as db_error:
        logger.error(f"Error retrieving documents from database: {db_error}")
        db.session.rollback()
        # Fall back to session documents
        if 'uploads' in session and session['uploads']:
            logger.info("Using session document data as fallback")
            flash("Using cached document data due to database connection issues.", "warning")
            return render_template('dashboard.html', documents=session['uploads'])
        else:
            flash("Unable to retrieve your documents. Please try again later.", "error")
    
    # Combine database documents with session documents
    # Session documents have priority to show the most up-to-date information
    session_doc_ids = [doc['id'] for doc in session['uploads']]
    
    # Filter out documents that are already in session
    for doc in db_documents:
        if doc.id not in session_doc_ids:
            # Add to session for future reference
            session_document = {
                'id': doc.id,
                'filename': doc.filename,
                'filetype': doc.filetype,
                'summary': doc.summary[:200] if doc.summary else "",
                'upload_time': doc.upload_time.strftime("%Y-%m-%d %H:%M:%S")
            }
            session['uploads'].append(session_document)
            session.modified = True
            
    # Return the template with documents
    return render_template('dashboard.html', documents=db_documents)

# Route to serve chat attachment files
@app.route('/attachments/<path:filename>')
@db_retry(max_retries=3)
def serve_attachment(filename):
    """Serve attachment files"""
    try:
        # Check if user is in session
        if 'session_id' not in session:
            flash("Please login to access attachments")
            return redirect(url_for('index'))
        
        # Check if the attachment exists in database and belongs to current session
        chat_message = ChatMessage.query.filter_by(
            attachment_filename=filename,
            session_id=session['session_id']
        ).first()
        
        if not chat_message:
            flash("Attachment not found or you don't have permission to access it")
            return redirect(url_for('index'))
        
        # Serve the file
        return send_from_directory(app.config['ATTACHMENT_FOLDER'], filename)
    except Exception as e:
        # Log the error
        logger.exception(f"Error serving attachment {filename}: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Return error message
        flash("An error occurred accessing this attachment. Please try again later.", "error")
        return redirect(url_for('dashboard'))

# API endpoint to generate flashcards
@app.route('/api/generate-flashcards/<int:document_id>', methods=['POST'])
@db_retry(max_retries=3)
def api_generate_flashcards(document_id):
    """Generate flashcards for a document"""
    try:
        # Check if user is in session
        if 'session_id' not in session:
            return jsonify({'error': 'Session expired. Please refresh the page.'}), 401
            
        # Retrieve document
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404
            
        # Check if user owns this document
        if document.session_id != session['session_id']:
            return jsonify({'error': 'Unauthorized access'}), 403
            
        # Check if document has content
        if not document.text_content:
            return jsonify({
                'error': 'This document has no extractable text content',
                'details': 'The file may be an image-based PDF or contain unsupported formatting.'
            }), 422
            
        # Generate flashcards
        flashcards, error = generate_flashcards(document.text_content, document_id)
        
        # Handle generation errors
        if error:
            return jsonify({
                'error': 'Failed to generate flashcards',
                'details': error
            }), 500
            
        # Return flashcards
        return jsonify({
            'success': True,
            'flashcards': flashcards,
            'count': len(flashcards)
        })
        
    except Exception as e:
        # Log the error
        logger.exception(f"Error generating flashcards: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Return error response
        return jsonify({
            'error': 'An unexpected error occurred',
            'details': str(e)
        }), 500

# Route to view flashcards for a document
@app.route('/flashcards/<int:document_id>')
@db_retry(max_retries=3)
def view_flashcards(document_id):
    """View flashcards for a document"""
    try:
        # Retrieve document
        document = Document.query.get_or_404(document_id)
        
        # Check if user owns this document
        if document.session_id != session.get('session_id'):
            flash("You don't have permission to view this document")
            return redirect(url_for('index'))
        
        # Check if flashcards exist in session
        doc_id_str = str(document_id)
        flashcards = []
        
        if 'flashcards' in session and doc_id_str in session['flashcards']:
            flashcards = session['flashcards'][doc_id_str]
            
        return render_template(
            'flashcards.html',
            document=document,
            flashcards=flashcards
        )
    except Exception as e:
        # Log the error
        logger.exception(f"Error viewing flashcards: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Show error message
        flash(f"An error occurred while retrieving flashcards: {str(e)}")
        return redirect(url_for('dashboard'))
    
# Route to view quiz for a document
@app.route('/quiz/<int:document_id>')
@db_retry(max_retries=3)
def view_quiz(document_id):
    """View quiz for a document"""
    try:
        # Retrieve document
        document = Document.query.get_or_404(document_id)
        
        # Check if user owns this document
        if document.session_id != session.get('session_id'):
            flash("You don't have permission to view this document")
            return redirect(url_for('index'))
        
        # Check if quiz exists in session
        doc_id_str = str(document_id)
        quiz_questions = []
        
        if 'quizzes' in session and doc_id_str in session['quizzes']:
            quiz_questions = session['quizzes'][doc_id_str]
            
        return render_template(
            'quiz.html',
            document=document,
            quiz_questions=quiz_questions
        )
    except Exception as e:
        # Log the error
        logger.exception(f"Error viewing quiz: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Show error message
        flash(f"An error occurred while retrieving quiz: {str(e)}")
        return redirect(url_for('dashboard'))

# Route to generate quiz for a document
@app.route('/api/quiz/<int:document_id>', methods=['POST'])
@db_retry(max_retries=3)
def api_generate_quiz(document_id):
    """Generate quiz for a document"""
    try:
        # Check if user is in session
        if 'session_id' not in session:
            return jsonify({'error': 'Session expired. Please refresh the page.'}), 401
            
        # Retrieve document
        document = Document.query.get_or_404(document_id)
        
        # Check if user owns this document
        if document.session_id != session.get('session_id'):
            return jsonify({'error': 'You do not have permission to access this document'}), 403
            
        # Check if document has text content
        if not document.text_content:
            return jsonify({'error': 'Document has no extractable content for quiz generation'}), 400
            
        # Generate quiz
        quiz_questions, error = generate_quiz(document.text_content, document_id)
        
        if error:
            return jsonify({'error': error}), 500
            
        return jsonify({
            'success': True,
            'quiz': quiz_questions
        })
        
    except Exception as e:
        # Log the error
        logger.exception(f"Error generating quiz: {e}")
        
        # Ensure database session is clean
        db.session.rollback()
        
        # Return error message
        return jsonify({'error': str(e)}), 500

# Route to submit quiz answers and get results
@app.route('/api/quiz/submit/<int:document_id>', methods=['POST'])
def submit_quiz(document_id):
    """Submit quiz answers and get results"""
    try:
        # Check if quiz exists in session
        doc_id_str = str(document_id)
        if 'quizzes' not in session or doc_id_str not in session['quizzes']:
            return jsonify({'error': 'Quiz not found. Please generate a quiz first.'}), 404
            
        # Get the quiz questions and user answers
        quiz_questions = session['quizzes'][doc_id_str]
        user_answers = request.json.get('answers', {})
        
        if not user_answers:
            return jsonify({'error': 'No answers submitted'}), 400
            
        # Calculate the score and prepare results
        results = []
        correct_count = 0
        total_count = len(quiz_questions)
        
        for i, question in enumerate(quiz_questions):
            question_idx = str(i)  # Using index as the question identifier
            
            user_choice = user_answers.get(question_idx)
            correct_choice = question['correct_answer']
            
            is_correct = user_choice == correct_choice
            if is_correct:
                correct_count += 1
                
            # Add result for this question
            results.append({
                'question': question['question'],
                'options': question['options'],
                'user_choice': user_choice,
                'correct_choice': correct_choice,
                'is_correct': is_correct
            })
            
        # Calculate percentage score
        score_percentage = int((correct_count / total_count) * 100) if total_count > 0 else 0
        
        # Prepare score data
        score = {
            'correct': correct_count,
            'total': total_count,
            'percentage': score_percentage
        }
        
        return jsonify({
            'success': True,
            'results': results,
            'score': score
        })
        
    except Exception as e:
        # Log the error
        logger.exception(f"Error submitting quiz: {e}")
        
        # Return error message
        return jsonify({'error': str(e)}), 500

# Route to clear session history
@app.route('/clear-history')
def clear_history():
    """Clear session history of uploads and chat"""
    if 'uploads' in session:
        session.pop('uploads')
    
    if 'chat_history' in session:
        session.pop('chat_history')
    
    if 'flashcards' in session:
        session.pop('flashcards')
        
    if 'quizzes' in session:
        session.pop('quizzes')
    
    # Make sure to persist changes
    session.modified = True
    
    flash("Your browsing history has been cleared.", "success")
    return redirect(url_for('dashboard'))

# Static file routes
@app.route('/static/<path:path>')
def serve_static(path):
    return send_from_directory('static', path)

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return render_template('error.html', error="Page not found"), 404

@app.errorhandler(500)
def server_error(error):
    return render_template('error.html', error="Internal server error"), 500

# Translation API endpoint
@app.route('/api/regenerate-tips/<int:document_id>', methods=['POST'])
@db_retry()
def regenerate_tips_api(document_id):
    """API endpoint to regenerate document interaction tips"""
    try:
        # Verify document exists
        document = Document.query.get(document_id)
        if not document:
            logger.error(f"Document not found: {document_id}")
            return jsonify({"error": "Document not found"}), 404
            
        # Verify user has access to this document
        if document.session_id != session.get('session_id'):
            logger.warning(f"Unauthorized access attempt to document: {document_id}")
            return jsonify({"error": "Unauthorized access"}), 403
            
        # Generate new tips based on document summary
        new_tips, error = generate_interaction_tips(document.summary, document.filetype)
        
        if error:
            logger.error(f"Error regenerating tips: {error}")
            return jsonify({"error": "Failed to generate new tips"}), 500
            
        # Update document with new tips
        document.interaction_tips = new_tips
        db.session.commit()
        logger.info(f"Tips regenerated for document {document_id}")
        
        # Also update session storage if this document is in session
        if 'uploads' in session:
            for i, doc in enumerate(session['uploads']):
                if str(doc.get('id')) == str(document_id):
                    session['uploads'][i]['interaction_tips'] = new_tips
                    session.modified = True
                    logger.debug(f"Updated tips in session for document {document_id}")
                    break
        
        # Return new tips
        return jsonify({"tips": new_tips}), 200
        
    except Exception as e:
        db.session.rollback()
        logger.exception(f"Error in regenerate_tips_api: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route('/api/translate/<int:document_id>', methods=['POST'])
@db_retry()
def translate_document_api(document_id):
    """Translate document content to the specified language"""
    try:
        # Get target language from request
        data = request.get_json()
        if not data or 'language' not in data:
            return jsonify({'error': 'Target language is required'}), 400
            
        target_language = data['language']
        
        # Validate target language
        if target_language not in TRANSLATION_LANGUAGES:
            return jsonify({'error': f'Unsupported target language: {target_language}'}), 400
            
        # Get document from database
        document = Document.query.get(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404
            
        # Check if document has text content
        if not document.text_content:
            return jsonify({'error': 'Document has no extractable text content'}), 400
            
        # Check if document is already translated to requested language
        if document.has_translation and document.translation_language == target_language:
            return jsonify({
                'message': 'Document already translated to this language',
                'translated_text': document.translated_text,
                'language_guide': document.language_guide,
                'language_info': TRANSLATION_LANGUAGES[target_language]
            }), 200
            
        # Translate the document
        logger.info(f"Translating document {document_id} to {target_language}")
        translated_text, language_guide, error = translate_document(document.text_content, target_language)
        
        if error:
            logger.error(f"Error translating document: {error}")
            return jsonify({'error': error}), 500
            
        # Save the translation to the database
        document.has_translation = True
        document.translated_text = translated_text
        document.translation_language = target_language
        document.language_guide = language_guide
        db.session.commit()
        
        # Return success response
        return jsonify({
            'message': 'Document successfully translated',
            'translated_text': translated_text,
            'language_guide': language_guide,
            'language_info': TRANSLATION_LANGUAGES[target_language]
        }), 200
        
    except Exception as e:
        logger.exception(f"Error in translation API: {e}")
        db.session.rollback()
        return jsonify({'error': f'Translation failed: {str(e)}'}), 500

# Route to view translated document
@app.route('/document/<int:document_id>/translate/<language>')
@db_retry()
def view_translated_document(document_id, language):
    """View document with translation in the specified language"""
    try:
        # Get document from database
        document = Document.query.get(document_id)
        if not document:
            flash("Document not found", "error")
            return redirect(url_for('dashboard'))
            
        # Check if this document has the requested translation
        has_requested_translation = document.has_translation and document.translation_language == language
        
        # If the document doesn't have a translation yet, or has a different language translation,
        # we'll redirect to the regular document view and let the client-side handle the translation request
        if not has_requested_translation:
            flash(f"Translation to {TRANSLATION_LANGUAGES.get(language, {}).get('name', language)} will be generated", "info")
            return redirect(url_for('view_document', document_id=document_id))
            
        # Get chat messages for this document
        chat_messages = ChatMessage.query.filter_by(document_id=document_id).order_by(ChatMessage.created_at).all()
        
        # Get session chat history as fallback (likely empty if we have database messages)
        session_chat_history = []
        if 'chat_history' in session and str(document_id) in session['chat_history']:
            session_chat_history = session['chat_history'][str(document_id)]
        
        # Render the document template with translation view
        return render_template(
            'document.html', 
            document=document,
            chat_messages=chat_messages,
            session_chat_history=session_chat_history,
            translation_languages=TRANSLATION_LANGUAGES,
            view_translation=True,
            language_code=language,
            language_info=TRANSLATION_LANGUAGES.get(language, {})
        )
        
    except Exception as e:
        logger.exception(f"Error viewing translated document: {e}")
        flash(f"Error loading translated document: {str(e)}", "error")
        return redirect(url_for('dashboard'))

# Run the application
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)